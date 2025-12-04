"""
Document upload and processing API endpoints.

Provides endpoints for:
- Uploading documents (PDF, DOCX, TXT) for an RFP
- Processing documents and adding to RAG index
- Listing and deleting uploaded documents
"""

import logging
import os
import sys
import uuid
from datetime import datetime, timezone
from pathlib import Path

from app.dependencies import RFPDep
from fastapi import APIRouter, BackgroundTasks, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel

# Add project root to path
project_root = os.path.dirname(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
)
if project_root not in sys.path:
    sys.path.insert(0, project_root)

logger = logging.getLogger(__name__)
router = APIRouter()

# Allowed file extensions
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


class UploadedDocument(BaseModel):
    """Response model for uploaded document."""

    id: str
    filename: str
    file_type: str
    file_size: int
    uploaded_at: str
    status: str = "pending"
    chunks_count: int | None = None
    error: str | None = None


class DocumentListResponse(BaseModel):
    """Response model for document list."""

    documents: list[UploadedDocument]
    total: int


class ProcessingStatus(BaseModel):
    """Status of document processing."""

    document_id: str
    status: str  # pending, processing, completed, failed
    progress: float = 0
    chunks_created: int = 0
    error: str | None = None


# In-memory storage for processing status (would be Redis in production)
_processing_status: dict[str, ProcessingStatus] = {}


def get_document_upload_dir(rfp_id: str) -> Path:
    """Get the upload directory for an RFP's documents."""
    upload_dir = Path("/app/data/uploads") / rfp_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def extract_text_from_file(filepath: Path) -> str:
    """Extract text content from a file with OCR fallback for scanned PDFs."""
    text = ""
    suffix = filepath.suffix.lower()

    try:
        if suffix == ".txt" or suffix == ".md":
            text = filepath.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            try:
                import fitz  # type: ignore[import-untyped] # PyMuPDF

                doc = fitz.open(str(filepath))
                for page in doc:
                    text += page.get_text()
                doc.close()
            except ImportError:
                logger.warning("PyMuPDF not installed, trying pdfplumber")
                try:
                    import pdfplumber  # type: ignore[import-untyped]

                    with pdfplumber.open(str(filepath)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except ImportError as err:
                    raise HTTPException(
                        status_code=500, detail="PDF processing libraries not available"
                    ) from err

            # If very little text extracted, try OCR (scanned PDF)
            if len(text.strip()) < 100:
                logger.info(f"PDF has little text ({len(text.strip())} chars), attempting OCR")
                try:
                    from pdf2image import convert_from_path
                    import pytesseract

                    # Convert PDF to images (limit to first 20 pages for performance)
                    images = convert_from_path(
                        str(filepath),
                        dpi=200,
                        first_page=1,
                        last_page=20
                    )

                    ocr_text = ""
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        ocr_text += f"\n[Page {i+1}]\n{page_text}"

                    # Use OCR text if we got more content
                    if len(ocr_text.strip()) > len(text.strip()):
                        logger.info(f"OCR extracted {len(ocr_text.strip())} chars vs {len(text.strip())} from standard extraction")
                        text = ocr_text

                except ImportError:
                    logger.warning("OCR libraries (pytesseract, pdf2image) not available for scanned PDF")
                except Exception as e:
                    logger.warning(f"OCR failed: {e}")
        elif suffix in {".docx", ".doc"}:
            try:
                from docx import Document

                doc = Document(str(filepath))
                text = "\n".join([para.text for para in doc.paragraphs])
            except ImportError as err:
                raise HTTPException(
                    status_code=500, detail="DOCX processing library not available"
                ) from err
        elif suffix in {".xlsx", ".xls"}:
            try:
                import pandas as pd

                df = pd.read_excel(str(filepath))
                text = df.to_string()
            except ImportError as err:
                raise HTTPException(
                    status_code=500, detail="Excel processing library not available"
                ) from err
    except Exception as e:
        logger.exception("Failed to extract text from %s", filepath)
        raise HTTPException(
            status_code=500, detail=f"Failed to extract text: {e!s}"
        ) from e

    return text


def process_document_for_rag(
    document_id: str,
    filepath: Path,
    rfp_id: str,
    filename: str,
):
    """Background task to process document and add to RAG index."""
    global _processing_status

    _processing_status[document_id] = ProcessingStatus(
        document_id=document_id,
        status="processing",
        progress=10,
    )

    try:
        # Extract text
        text = extract_text_from_file(filepath)
        _processing_status[document_id].progress = 40

        if not text.strip():
            _processing_status[document_id].status = "failed"
            _processing_status[document_id].error = "No text content extracted"
            return

        # Chunk the text
        chunk_size = 512
        overlap = 50
        chunks = []
        words = text.split()

        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i : i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)

        _processing_status[document_id].progress = 60
        _processing_status[document_id].chunks_created = len(chunks)

        # Add to RAG index
        try:
            from src.rag.chroma_rag_engine import get_rag_engine

            rag_engine = get_rag_engine()
            if not rag_engine:
                raise RuntimeError("RAG engine not available")

            # Create document IDs and metadata for each chunk
            doc_ids = [f"{document_id}_chunk_{i}" for i in range(len(chunks))]
            docs_for_chroma = [
                {
                    "content": chunk,
                    "source": filename,
                    "rfp_id": rfp_id,
                    "document_id": document_id,
                    "chunk_index": str(i),
                    "upload_date": datetime.now(timezone.utc).isoformat(),
                }
                for i, chunk in enumerate(chunks)
            ]

            added = rag_engine.add_documents(
                documents=docs_for_chroma,
                ids=doc_ids,
            )

            _processing_status[document_id].progress = 100
            _processing_status[document_id].status = "completed"
            _processing_status[document_id].chunks_created = added
            logger.info(f"Added {added} chunks from {filename} to RAG index")

        except Exception as e:
            logger.error(f"Failed to add document to RAG: {e}")
            _processing_status[document_id].status = "completed"
            _processing_status[document_id].progress = 100
            # Still mark as completed since file is saved

    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        _processing_status[document_id].status = "failed"
        _processing_status[document_id].error = str(e)


@router.post("/{rfp_id}/upload", response_model=UploadedDocument)
async def upload_document(
    rfp: RFPDep,
    file: UploadFile = File(...),  # noqa: B008
    background_tasks: BackgroundTasks = None,
):
    """
    Upload a document for an RFP.

    Supports PDF, DOCX, DOC, TXT, MD, XLSX, XLS files up to 50MB.
    Documents are processed in the background and added to the RAG index.
    """
    # Validate file extension
    filename = file.filename or "document"
    file_ext = Path(filename).suffix.lower()

    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Supported: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read file content
    content = await file.read()
    file_size = len(content)

    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024 * 1024)}MB",
        )

    if file_size == 0:
        raise HTTPException(status_code=400, detail="Empty file")

    # Generate document ID and save file
    document_id = f"doc-{uuid.uuid4().hex[:12]}"
    upload_dir = get_document_upload_dir(rfp.rfp_id)
    filepath = upload_dir / f"{document_id}{file_ext}"

    filepath.write_bytes(content)

    # Schedule background processing
    if background_tasks:
        background_tasks.add_task(
            process_document_for_rag,
            document_id,
            filepath,
            rfp.rfp_id,
            filename,
        )

    return UploadedDocument(
        id=document_id,
        filename=filename,
        file_type=file_ext.lstrip("."),
        file_size=file_size,
        uploaded_at=datetime.now(timezone.utc).isoformat(),
        status="processing",
    )


@router.get("/{rfp_id}/uploads", response_model=DocumentListResponse)
async def list_uploaded_documents(rfp: RFPDep):
    """List all uploaded documents for an RFP."""
    upload_dir = get_document_upload_dir(rfp.rfp_id)
    documents = []

    if upload_dir.exists():
        for filepath in upload_dir.iterdir():
            if filepath.is_file():
                doc_id = filepath.stem
                status = _processing_status.get(doc_id)

                documents.append(
                    UploadedDocument(
                        id=doc_id,
                        filename=filepath.name,
                        file_type=filepath.suffix.lstrip("."),
                        file_size=filepath.stat().st_size,
                        uploaded_at=datetime.fromtimestamp(
                            filepath.stat().st_mtime, tz=timezone.utc
                        ).isoformat(),
                        status=status.status if status else "completed",
                        chunks_count=status.chunks_created if status else None,
                        error=status.error if status else None,
                    )
                )

    return DocumentListResponse(
        documents=sorted(documents, key=lambda d: d.uploaded_at, reverse=True),
        total=len(documents),
    )


@router.get("/{rfp_id}/uploads/{document_id}/status", response_model=ProcessingStatus)
async def get_processing_status(rfp: RFPDep, document_id: str):
    """Get the processing status of an uploaded document."""
    if document_id not in _processing_status:
        # Check if file exists
        upload_dir = get_document_upload_dir(rfp.rfp_id)
        matching = list(upload_dir.glob(f"{document_id}*"))

        if matching:
            return ProcessingStatus(
                document_id=document_id,
                status="completed",
                progress=100,
            )

        raise HTTPException(status_code=404, detail="Document not found")

    return _processing_status[document_id]


@router.delete("/{rfp_id}/uploads/{document_id}")
async def delete_uploaded_document(rfp: RFPDep, document_id: str):
    """Delete an uploaded document."""
    upload_dir = get_document_upload_dir(rfp.rfp_id)
    matching = list(upload_dir.glob(f"{document_id}*"))

    if not matching:
        raise HTTPException(status_code=404, detail="Document not found")

    for filepath in matching:
        filepath.unlink()

    # Clean up processing status
    if document_id in _processing_status:
        del _processing_status[document_id]

    return {"status": "deleted", "document_id": document_id}


@router.get("/{rfp_id}/uploads/{document_id}/content")
async def get_document_content(
    rfp: RFPDep,
    document_id: str,
):
    """Get extracted text content from an uploaded document."""
    upload_dir = get_document_upload_dir(rfp.rfp_id)

    # Find the document file
    matching_files = list(upload_dir.glob(f"{document_id}.*"))
    if not matching_files:
        raise HTTPException(status_code=404, detail="Document not found")

    filepath = matching_files[0]

    try:
        text = extract_text_from_file(filepath)
        return {
            "document_id": document_id,
            "filename": filepath.name,
            "content": text,
            "char_count": len(text),
            "word_count": len(text.split()),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract content: {e!s}")


@router.get("/{rfp_id}/uploads/{document_id}/download")
async def download_document(
    rfp: RFPDep,
    document_id: str,
):
    """Download an uploaded document."""
    upload_dir = get_document_upload_dir(rfp.rfp_id)

    # Find the document file
    matching_files = list(upload_dir.glob(f"{document_id}.*"))
    if not matching_files:
        raise HTTPException(status_code=404, detail="Document not found")

    filepath = matching_files[0]

    # Use the actual filename (without the document_id prefix)
    original_filename = filepath.name

    return FileResponse(
        path=str(filepath),
        filename=original_filename,
        media_type="application/octet-stream"
    )
