"""
Document reader utility for extracting text from RFP attachments.

Supports PDF and DOCX files for use in AI proposal generation.
"""

import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Try to import document parsing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available - PDF reading disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX reading disabled")


def extract_text_from_pdf(file_path: str, max_pages: int = 50) -> Optional[str]:
    """
    Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file
        max_pages: Maximum number of pages to extract (default 50)

    Returns:
        Extracted text content or None if extraction fails
    """
    if not PDF_AVAILABLE:
        logger.warning("PDF reading not available - PyPDF2 not installed")
        return None

    if not os.path.exists(file_path):
        logger.warning(f"PDF file not found: {file_path}")
        return None

    try:
        reader = PdfReader(file_path)
        text_parts = []

        num_pages = min(len(reader.pages), max_pages)
        for i in range(num_pages):
            page = reader.pages[i]
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        if text_parts:
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} chars from PDF ({num_pages} pages): {file_path}")
            return full_text.strip()

        logger.warning(f"No text extracted from PDF: {file_path}")
        return None

    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {e}")
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content or None if extraction fails
    """
    if not DOCX_AVAILABLE:
        logger.warning("DOCX reading not available - python-docx not installed")
        return None

    if not os.path.exists(file_path):
        logger.warning(f"DOCX file not found: {file_path}")
        return None

    try:
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        if text_parts:
            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} chars from DOCX: {file_path}")
            return full_text.strip()

        logger.warning(f"No text extracted from DOCX: {file_path}")
        return None

    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}")
        return None


def extract_text_from_document(file_path: str) -> Optional[str]:
    """
    Extract text from a document based on its file extension.

    Supports PDF and DOCX files.

    Args:
        file_path: Path to the document file

    Returns:
        Extracted text content or None if extraction fails
    """
    if not file_path:
        return None

    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif extension == ".txt":
        # Plain text files
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            logger.info(f"Read {len(content)} chars from text file: {file_path}")
            return content.strip()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return None
    else:
        logger.warning(f"Unsupported document format: {extension} for {file_path}")
        return None


def extract_all_document_content(
    documents: list[dict],
    max_total_chars: int = 100000
) -> dict:
    """
    Extract text content from all available documents.

    Args:
        documents: List of document dicts with 'file_path' and 'filename' keys
        max_total_chars: Maximum total characters to extract (default 100k)

    Returns:
        Dictionary with 'documents' list and 'total_chars' count
    """
    extracted = []
    total_chars = 0

    for doc in documents:
        file_path = doc.get("file_path")
        filename = doc.get("filename", "Unknown")
        doc_type = doc.get("document_type", "attachment")

        if not file_path:
            logger.debug(f"No file_path for document: {filename}")
            continue

        if total_chars >= max_total_chars:
            logger.warning(f"Reached max chars limit ({max_total_chars}), skipping remaining documents")
            break

        content = extract_text_from_document(file_path)

        if content:
            # Truncate if needed to stay within limit
            remaining_chars = max_total_chars - total_chars
            if len(content) > remaining_chars:
                content = content[:remaining_chars] + "\n[... content truncated ...]"

            extracted.append({
                "filename": filename,
                "document_type": doc_type,
                "content": content,
                "char_count": len(content)
            })
            total_chars += len(content)

    logger.info(f"Extracted content from {len(extracted)} documents ({total_chars} total chars)")

    return {
        "documents": extracted,
        "total_chars": total_chars,
        "document_count": len(extracted)
    }
