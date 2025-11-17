"""
Document processor for bid submission format conversion.
"""
import os
import logging
from typing import Dict, List, Optional
from io import BytesIO
import base64

logger = logging.getLogger(__name__)

# Try to import PDF/DOCX libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available - PDF generation disabled")

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX generation disabled")


class DocumentProcessor:
    """Process bid documents for submission."""

    def __init__(self):
        """Initialize document processor."""
        self.supported_formats = []

        if REPORTLAB_AVAILABLE:
            self.supported_formats.append("PDF")
        if DOCX_AVAILABLE:
            self.supported_formats.append("DOCX")

        self.supported_formats.append("HTML")  # Always available
        self.supported_formats.append("JSON")  # Always available

        logger.info(f"Document processor initialized. Supported formats: {self.supported_formats}")

    def convert_format(
        self,
        bid_document: Dict,
        target_format: str
    ) -> bytes:
        """
        Convert bid document to target format.

        Args:
            bid_document: Bid document data
            target_format: Target format (PDF, DOCX, HTML, JSON)

        Returns:
            Document content as bytes
        """
        target_format = target_format.upper()

        if target_format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {target_format}")

        if target_format == "PDF":
            return self._convert_to_pdf(bid_document)
        elif target_format == "DOCX":
            return self._convert_to_docx(bid_document)
        elif target_format == "HTML":
            return self._convert_to_html(bid_document)
        elif target_format == "JSON":
            return self._convert_to_json(bid_document)
        else:
            raise ValueError(f"Format not implemented: {target_format}")

    def _convert_to_pdf(self, bid_document: Dict) -> bytes:
        """Convert to PDF format."""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("PDF conversion not available - reportlab not installed")

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title = Paragraph(
            f"<b>{bid_document.get('title', 'Bid Document')}</b>",
            styles['Title']
        )
        story.append(title)
        story.append(Spacer(1, 12))

        # Content sections
        content = bid_document.get('content_markdown', bid_document.get('content', ''))

        # Simple paragraph rendering (in production, use more sophisticated markdown parsing)
        for line in content.split('\n\n'):
            if line.strip():
                para = Paragraph(line.strip(), styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))

        # Build PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        logger.info(f"Generated PDF: {len(pdf_bytes)} bytes")
        return pdf_bytes

    def _convert_to_docx(self, bid_document: Dict) -> bytes:
        """Convert to DOCX format."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX conversion not available - python-docx not installed")

        document = Document()

        # Add title
        document.add_heading(bid_document.get('title', 'Bid Document'), 0)

        # Add content
        content = bid_document.get('content_markdown', bid_document.get('content', ''))

        for line in content.split('\n\n'):
            if line.strip():
                document.add_paragraph(line.strip())

        # Save to buffer
        buffer = BytesIO()
        document.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        logger.info(f"Generated DOCX: {len(docx_bytes)} bytes")
        return docx_bytes

    def _convert_to_html(self, bid_document: Dict) -> bytes:
        """Convert to HTML format."""
        html_content = bid_document.get('content_html', '')

        if not html_content:
            # Generate from markdown or plain content
            content = bid_document.get('content_markdown', bid_document.get('content', ''))
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{bid_document.get('title', 'Bid Document')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; }}
        p {{ line-height: 1.6; }}
    </style>
</head>
<body>
    <h1>{bid_document.get('title', 'Bid Document')}</h1>
    <div>{content.replace(chr(10), '<br>')}</div>
</body>
</html>
"""

        return html_content.encode('utf-8')

    def _convert_to_json(self, bid_document: Dict) -> bytes:
        """Convert to JSON format."""
        import json
        json_str = json.dumps(bid_document, indent=2, default=str)
        return json_str.encode('utf-8')

    def assemble_package(
        self,
        bid_document: Dict,
        portal_requirements: Dict
    ) -> Dict:
        """
        Assemble complete submission package.

        Args:
            bid_document: Bid document data
            portal_requirements: Portal-specific requirements

        Returns:
            Complete submission package
        """
        package = {
            "primary_document": bid_document,
            "attachments": [],
            "forms": {},
            "certifications": []
        }

        # Convert to required format
        required_format = portal_requirements.get("format", "PDF")
        primary_content = self.convert_format(bid_document, required_format)

        package["primary_document"]["content_bytes"] = base64.b64encode(primary_content).decode('utf-8')
        package["primary_document"]["format"] = required_format
        package["primary_document"]["file_size"] = len(primary_content)

        # Add required forms
        required_forms = portal_requirements.get("required_forms", [])
        for form_name in required_forms:
            package["forms"][form_name] = self._generate_standard_form(form_name, bid_document)

        # Add certifications
        required_certs = portal_requirements.get("required_certifications", [])
        for cert_name in required_certs:
            package["certifications"].append({
                "name": cert_name,
                "status": "included",
                "document_id": f"{cert_name}_cert"
            })

        logger.info(f"Assembled submission package with {len(package['attachments'])} attachments")
        return package

    def _generate_standard_form(self, form_name: str, bid_document: Dict) -> Dict:
        """Generate standard government form."""
        # Simplified form generation
        # In production, this would use actual form templates
        form_data = {
            "form_name": form_name,
            "vendor_info": {
                "name": bid_document.get("vendor_name", ""),
                "cage_code": bid_document.get("cage_code", ""),
                "duns": bid_document.get("duns_number", "")
            },
            "generated_at": str(os.times())
        }

        return form_data

    def validate_package(
        self,
        package: Dict,
        requirements: Dict
    ) -> List[str]:
        """
        Validate submission package.

        Args:
            package: Submission package
            requirements: Portal requirements

        Returns:
            List of validation errors (empty if valid)
        """
        errors = []

        # Check file size
        max_size = requirements.get("max_file_size", 100 * 1024 * 1024)
        if package["primary_document"].get("file_size", 0) > max_size:
            errors.append(f"File size exceeds limit of {max_size} bytes")

        # Check required forms
        required_forms = requirements.get("required_forms", [])
        for form_name in required_forms:
            if form_name not in package["forms"]:
                errors.append(f"Missing required form: {form_name}")

        # Check required certifications
        required_certs = requirements.get("required_certifications", [])
        cert_names = [cert["name"] for cert in package["certifications"]]
        for cert_name in required_certs:
            if cert_name not in cert_names:
                errors.append(f"Missing required certification: {cert_name}")

        return errors
